"""Unit tests untuk validasi CLI dan model data convert_md_to_docx."""

import sys
import os
import pytest

# Tambahkan path skrip agar bisa import modul
sys.path.insert(0, os.path.dirname(__file__))

from convert_md_to_docx import (
    TextRun,
    Paragraph,
    BulletItem,
    ChecklistItem,
    CodeBlock,
    Table,
    Section,
    main,
    parse_inline,
    parse_markdown,
)


# ---------------------------------------------------------------------------
# CLI Tests
# ---------------------------------------------------------------------------

class TestCLI:
    """Test CLI entry point."""

    def test_main_tanpa_argumen_menampilkan_usage(self, monkeypatch):
        """main() tanpa argumen harus keluar dengan SystemExit (argparse error)."""
        monkeypatch.setattr("sys.argv", ["convert_md_to_docx.py"])
        with pytest.raises(SystemExit) as exc_info:
            main()
        # argparse exits with code 2 for missing required arguments
        assert exc_info.value.code == 2


# ---------------------------------------------------------------------------
# Dataclass Default Value Tests
# ---------------------------------------------------------------------------

class TestTextRun:
    """Test TextRun dataclass instantiation dan default values."""

    def test_default_bold(self):
        run = TextRun(text="hello")
        assert run.bold is False

    def test_default_italic(self):
        run = TextRun(text="hello")
        assert run.italic is False

    def test_default_is_risk_indicator(self):
        run = TextRun(text="hello")
        assert run.is_risk_indicator is False

    def test_default_font_name(self):
        run = TextRun(text="hello")
        assert run.font_name is None

    def test_default_font_size(self):
        run = TextRun(text="hello")
        assert run.font_size is None

    def test_default_color(self):
        run = TextRun(text="hello")
        assert run.color is None

    def test_text_value(self):
        run = TextRun(text="test text")
        assert run.text == "test text"


class TestParagraph:
    """Test Paragraph dataclass instantiation dan default values."""

    def test_default_runs_empty(self):
        para = Paragraph()
        assert para.runs == []

    def test_runs_is_list(self):
        para = Paragraph()
        assert isinstance(para.runs, list)


class TestBulletItem:
    """Test BulletItem dataclass instantiation dan default values."""

    def test_default_runs_empty(self):
        item = BulletItem()
        assert item.runs == []

    def test_default_level(self):
        item = BulletItem()
        assert item.level == 0


class TestChecklistItem:
    """Test ChecklistItem dataclass instantiation dan default values."""

    def test_default_runs_empty(self):
        item = ChecklistItem()
        assert item.runs == []

    def test_default_checked(self):
        item = ChecklistItem()
        assert item.checked is False


class TestCodeBlock:
    """Test CodeBlock dataclass instantiation dan default values."""

    def test_default_content(self):
        block = CodeBlock()
        assert block.content == ""


class TestTable:
    """Test Table dataclass instantiation dan default values."""

    def test_default_headers_empty(self):
        table = Table()
        assert table.headers == []

    def test_default_rows_empty(self):
        table = Table()
        assert table.rows == []

    def test_headers_is_list(self):
        table = Table()
        assert isinstance(table.headers, list)

    def test_rows_is_list(self):
        table = Table()
        assert isinstance(table.rows, list)


class TestSection:
    """Test Section dataclass instantiation dan default values."""

    def test_default_level(self):
        section = Section()
        assert section.level == 1

    def test_default_text(self):
        section = Section()
        assert section.text == ""

    def test_default_children_empty(self):
        section = Section()
        assert section.children == []

    def test_children_is_list(self):
        section = Section()
        assert isinstance(section.children, list)


# ---------------------------------------------------------------------------
# parse_inline Tests
# ---------------------------------------------------------------------------

class TestParseInline:
    """Test parse_inline helper function."""

    def test_plain_text(self):
        runs = parse_inline("hello world")
        assert len(runs) == 1
        assert runs[0].text == "hello world"
        assert runs[0].bold is False
        assert runs[0].italic is False

    def test_bold_text(self):
        runs = parse_inline("this is **bold** text")
        assert len(runs) == 3
        assert runs[0].text == "this is "
        assert runs[1].text == "bold"
        assert runs[1].bold is True
        assert runs[2].text == " text"

    def test_italic_text(self):
        runs = parse_inline("this is *italic* text")
        assert len(runs) == 3
        assert runs[0].text == "this is "
        assert runs[1].text == "italic"
        assert runs[1].italic is True
        assert runs[2].text == " text"

    def test_risk_emoji_red(self):
        runs = parse_inline("Risk: 🔴 Critical")
        assert any(r.is_risk_indicator and r.text == "🔴" for r in runs)

    def test_risk_emoji_orange(self):
        runs = parse_inline("🟠 High")
        assert runs[0].text == "🟠"
        assert runs[0].is_risk_indicator is True

    def test_risk_emoji_yellow(self):
        runs = parse_inline("🟡 Medium")
        assert runs[0].text == "🟡"
        assert runs[0].is_risk_indicator is True

    def test_risk_emoji_green(self):
        runs = parse_inline("🟢 Low")
        assert runs[0].text == "🟢"
        assert runs[0].is_risk_indicator is True

    def test_mixed_bold_and_italic(self):
        runs = parse_inline("**bold** and *italic*")
        bold_runs = [r for r in runs if r.bold]
        italic_runs = [r for r in runs if r.italic]
        assert len(bold_runs) == 1
        assert bold_runs[0].text == "bold"
        assert len(italic_runs) == 1
        assert italic_runs[0].text == "italic"

    def test_empty_string(self):
        runs = parse_inline("")
        assert runs == []


# ---------------------------------------------------------------------------
# parse_markdown Tests
# ---------------------------------------------------------------------------

class TestParseMarkdownHeadings:
    """Test parsing heading H1, H2, H3."""

    def test_h1(self):
        sections = parse_markdown("# Title")
        assert len(sections) == 1
        assert sections[0].level == 1
        assert sections[0].text == "Title"

    def test_h2(self):
        sections = parse_markdown("## Subtitle")
        assert len(sections) == 1
        assert sections[0].level == 2
        assert sections[0].text == "Subtitle"

    def test_h3(self):
        sections = parse_markdown("### Sub-subtitle")
        assert len(sections) == 1
        assert sections[0].level == 3
        assert sections[0].text == "Sub-subtitle"

    def test_multiple_headings_preserve_order(self):
        md = "# First\n## Second\n### Third\n## Fourth"
        sections = parse_markdown(md)
        assert len(sections) == 4
        assert [s.level for s in sections] == [1, 2, 3, 2]
        assert [s.text for s in sections] == ["First", "Second", "Third", "Fourth"]


class TestParseMarkdownTable:
    """Test parsing tabel Markdown."""

    def test_simple_table(self):
        md = "# Test\n| A | B |\n|---|---|\n| 1 | 2 |\n| 3 | 4 |"
        sections = parse_markdown(md)
        tables = [c for c in sections[0].children if isinstance(c, Table)]
        assert len(tables) == 1
        assert tables[0].headers == ["A", "B"]
        assert tables[0].rows == [["1", "2"], ["3", "4"]]

    def test_table_with_emoji(self):
        md = "# Test\n| Risk | Level |\n|------|-------|\n| Data | 🔴 Critical |"
        sections = parse_markdown(md)
        tables = [c for c in sections[0].children if isinstance(c, Table)]
        assert len(tables) == 1
        assert "🔴" in tables[0].rows[0][1]


class TestParseMarkdownBulletList:
    """Test parsing bullet list with nested levels."""

    def test_simple_bullet(self):
        md = "# Test\n- item one\n- item two"
        sections = parse_markdown(md)
        bullets = [c for c in sections[0].children if isinstance(c, BulletItem)]
        assert len(bullets) == 2
        assert bullets[0].level == 0
        assert bullets[0].runs[0].text == "item one"

    def test_nested_bullet(self):
        md = "# Test\n- parent\n  - child\n    - grandchild"
        sections = parse_markdown(md)
        bullets = [c for c in sections[0].children if isinstance(c, BulletItem)]
        assert len(bullets) == 3
        assert bullets[0].level == 0
        assert bullets[1].level == 1
        assert bullets[2].level == 2


class TestParseMarkdownChecklist:
    """Test parsing checklist checked dan unchecked."""

    def test_unchecked(self):
        md = "# Test\n- [ ] unchecked item"
        sections = parse_markdown(md)
        items = [c for c in sections[0].children if isinstance(c, ChecklistItem)]
        assert len(items) == 1
        assert items[0].checked is False
        assert items[0].runs[0].text == "unchecked item"

    def test_checked(self):
        md = "# Test\n- [x] checked item"
        sections = parse_markdown(md)
        items = [c for c in sections[0].children if isinstance(c, ChecklistItem)]
        assert len(items) == 1
        assert items[0].checked is True


class TestParseMarkdownCodeBlock:
    """Test parsing code block."""

    def test_code_block(self):
        md = "# Test\n```\nprint('hello')\nprint('world')\n```"
        sections = parse_markdown(md)
        blocks = [c for c in sections[0].children if isinstance(c, CodeBlock)]
        assert len(blocks) == 1
        assert "print('hello')" in blocks[0].content
        assert "print('world')" in blocks[0].content

    def test_code_block_with_language(self):
        md = "# Test\n```python\nx = 1\n```"
        sections = parse_markdown(md)
        blocks = [c for c in sections[0].children if isinstance(c, CodeBlock)]
        assert len(blocks) == 1
        assert "x = 1" in blocks[0].content


class TestParseMarkdownInlineFormatting:
    """Test parsing bold dan italic inline in paragraphs."""

    def test_bold_in_paragraph(self):
        md = "# Test\nThis is **bold** text"
        sections = parse_markdown(md)
        paras = [c for c in sections[0].children if isinstance(c, Paragraph)]
        assert len(paras) == 1
        bold_runs = [r for r in paras[0].runs if r.bold]
        assert len(bold_runs) == 1
        assert bold_runs[0].text == "bold"

    def test_italic_in_paragraph(self):
        md = "# Test\nThis is *italic* text"
        sections = parse_markdown(md)
        paras = [c for c in sections[0].children if isinstance(c, Paragraph)]
        assert len(paras) == 1
        italic_runs = [r for r in paras[0].runs if r.italic]
        assert len(italic_runs) == 1
        assert italic_runs[0].text == "italic"


class TestParseMarkdownRiskEmoji:
    """Test parsing emoji indikator risiko."""

    def test_risk_emoji_in_text(self):
        md = "# Test\n🔴 Critical risk found"
        sections = parse_markdown(md)
        paras = [c for c in sections[0].children if isinstance(c, Paragraph)]
        risk_runs = [r for p in paras for r in p.runs if r.is_risk_indicator]
        assert len(risk_runs) >= 1
        assert risk_runs[0].text == "🔴"

    def test_all_risk_emojis(self):
        md = "# Test\n🔴 🟠 🟡 🟢"
        sections = parse_markdown(md)
        paras = [c for c in sections[0].children if isinstance(c, Paragraph)]
        risk_runs = [r for p in paras for r in p.runs if r.is_risk_indicator]
        assert len(risk_runs) == 4
        emojis = {r.text for r in risk_runs}
        assert emojis == {"🔴", "🟠", "🟡", "🟢"}


class TestParseMarkdownEdgeCases:
    """Test edge cases."""

    def test_empty_content(self):
        sections = parse_markdown("")
        assert sections == []

    def test_no_heading(self):
        md = "Just some text without heading"
        sections = parse_markdown(md)
        assert len(sections) == 1
        assert sections[0].level == 0

    def test_horizontal_rule(self):
        md = "# Test\n---\nSome text"
        sections = parse_markdown(md)
        children = sections[0].children
        # HR should be a Paragraph with "---"
        hr_paras = [c for c in children if isinstance(c, Paragraph) and any(r.text == "---" for r in c.runs)]
        assert len(hr_paras) == 1

    def test_content_before_first_heading(self):
        md = "Preamble text\n# Heading"
        sections = parse_markdown(md)
        assert len(sections) == 2
        assert sections[0].level == 0
        assert sections[1].level == 1


# ---------------------------------------------------------------------------
# Property-Based Tests (hypothesis)
# ---------------------------------------------------------------------------

from hypothesis import given, settings, strategies as st


class TestPropertyPreservasiUrutanSection:
    """Feature: sow-analysis-docx-conversion, Property 3: Preservasi Urutan Section

    For any file Markdown dengan beberapa section (heading), urutan kemunculan
    section dalam output parse_markdown harus sama persis dengan urutan dalam
    file Markdown sumber.

    **Validates: Requirements 7.2**
    """

    @settings(max_examples=100)
    @given(
        headings=st.lists(
            st.tuples(
                st.integers(min_value=1, max_value=3),
                st.text(
                    alphabet=st.characters(
                        whitelist_categories=("L", "N", "Zs"),
                        blacklist_characters="\n\r\x00#*|`-[]",
                    ),
                    min_size=1,
                    max_size=30,
                ),
            ),
            min_size=1,
            max_size=15,
        )
    )
    def test_urutan_heading_dipertahankan(self, headings):
        """Generate random MD dengan multiple heading, verifikasi urutan heading
        di output sama dengan input.

        **Validates: Requirements 7.2**
        """
        # Build Markdown string from generated headings
        md_lines = []
        for level, text in headings:
            # Ensure text is non-empty after stripping
            clean_text = text.strip()
            if not clean_text:
                clean_text = "heading"
            prefix = "#" * level
            md_lines.append(f"{prefix} {clean_text}")

        md_content = "\n".join(md_lines)

        # Parse the Markdown
        sections = parse_markdown(md_content)

        # Extract heading texts and levels from parsed output
        parsed_texts = [s.text for s in sections]
        parsed_levels = [s.level for s in sections]

        # Build expected texts and levels
        expected_texts = []
        expected_levels = []
        for level, text in headings:
            clean_text = text.strip()
            if not clean_text:
                clean_text = "heading"
            expected_texts.append(clean_text)
            expected_levels.append(level)

        # Verify order and count match
        assert len(sections) == len(headings), (
            f"Expected {len(headings)} sections, got {len(sections)}"
        )
        assert parsed_texts == expected_texts, (
            f"Heading texts order mismatch.\n"
            f"Expected: {expected_texts}\n"
            f"Got:      {parsed_texts}"
        )
        assert parsed_levels == expected_levels, (
            f"Heading levels order mismatch.\n"
            f"Expected: {expected_levels}\n"
            f"Got:      {parsed_levels}"
        )


class TestPropertyPemetaanHierarkiHeading:
    """Feature: sow-analysis-docx-conversion, Property 4: Pemetaan Hierarki Heading

    For any file Markdown yang mengandung heading level 1, 2, dan 3, setiap
    heading harus dipetakan ke heading Word dengan level hierarki yang sesuai
    (H1→Heading 1, H2→Heading 2, H3→Heading 3).

    **Validates: Requirements 2.1**
    """

    @settings(max_examples=100)
    @given(
        headings=st.lists(
            st.tuples(
                st.integers(min_value=1, max_value=3),
                st.text(
                    alphabet=st.characters(
                        whitelist_categories=("L", "N", "Zs"),
                        blacklist_characters="\n\r\x00#*|`-[]",
                    ),
                    min_size=1,
                    max_size=30,
                ),
            ),
            min_size=1,
            max_size=20,
        )
    )
    def test_heading_level_dipetakan_dengan_benar(self, headings):
        """Generate random heading levels (1-3), verifikasi setiap heading
        dipetakan ke level yang benar di output parse_markdown().

        **Validates: Requirements 2.1**
        """
        # Build Markdown string from generated headings
        md_lines = []
        expected_levels = []
        for level, text in headings:
            clean_text = text.strip()
            if not clean_text:
                clean_text = "heading"
            prefix = "#" * level
            md_lines.append(f"{prefix} {clean_text}")
            expected_levels.append(level)

        md_content = "\n".join(md_lines)

        # Parse the Markdown
        sections = parse_markdown(md_content)

        # Verify count matches
        assert len(sections) == len(headings), (
            f"Expected {len(headings)} sections, got {len(sections)}"
        )

        # Verify each heading level is correctly mapped
        for idx, (section, expected_level) in enumerate(zip(sections, expected_levels)):
            assert section.level == expected_level, (
                f"Heading #{idx}: expected level {expected_level}, "
                f"got level {section.level} for text '{section.text}'"
            )


class TestPropertyKonversiTabelFidelitasData:
    """Feature: sow-analysis-docx-conversion, Property 5: Konversi Tabel dengan Fidelitas Data

    For any tabel Markdown dengan header dan baris data, output parse_markdown()
    harus mengandung Table dengan jumlah kolom dan baris yang sama, serta seluruh
    data sel yang identik.

    **Validates: Requirements 2.2, 7.3**
    """

    @settings(max_examples=100)
    @given(
        num_cols=st.integers(min_value=1, max_value=5),
        num_rows=st.integers(min_value=1, max_value=5),
        data=st.data(),
    )
    def test_tabel_fidelitas_data(self, num_cols, num_rows, data):
        """Generate random tabel (random cols/rows/data), verifikasi jumlah kolom,
        baris, dan data sel identik setelah parsing.

        **Validates: Requirements 2.2, 7.3**
        """
        cell_strategy = st.text(
            alphabet=st.characters(
                whitelist_categories=("L", "N"),
                min_codepoint=65,
                max_codepoint=122,
            ),
            min_size=1,
            max_size=10,
        )

        # Generate header cells
        headers = [data.draw(cell_strategy) for _ in range(num_cols)]

        # Generate data rows
        row_data = []
        for _ in range(num_rows):
            row = [data.draw(cell_strategy) for _ in range(num_cols)]
            row_data.append(row)

        # Build Markdown table string
        header_line = "| " + " | ".join(headers) + " |"
        separator_line = "| " + " | ".join(["---"] * num_cols) + " |"
        data_lines = []
        for row in row_data:
            data_lines.append("| " + " | ".join(row) + " |")

        md_content = "# Test\n" + header_line + "\n" + separator_line + "\n" + "\n".join(data_lines)

        # Parse the Markdown
        sections = parse_markdown(md_content)

        # Find the Table object in parsed output
        tables = []
        for section in sections:
            for child in section.children:
                if isinstance(child, Table):
                    tables.append(child)

        assert len(tables) == 1, f"Expected 1 table, got {len(tables)}"

        table = tables[0]

        # Verify number of headers matches columns
        assert len(table.headers) == num_cols, (
            f"Expected {num_cols} headers, got {len(table.headers)}"
        )

        # Verify number of rows matches
        assert len(table.rows) == num_rows, (
            f"Expected {num_rows} rows, got {len(table.rows)}"
        )

        # Verify header data is identical
        assert table.headers == headers, (
            f"Header mismatch.\nExpected: {headers}\nGot: {table.headers}"
        )

        # Verify each cell data is identical
        for row_idx, (expected_row, actual_row) in enumerate(zip(row_data, table.rows)):
            assert len(actual_row) == num_cols, (
                f"Row {row_idx}: expected {num_cols} cells, got {len(actual_row)}"
            )
            assert actual_row == expected_row, (
                f"Row {row_idx} data mismatch.\n"
                f"Expected: {expected_row}\n"
                f"Got: {actual_row}"
            )


import tempfile
from docx import Document as DocxDocument

from convert_md_to_docx import build_docx


class TestPropertyPreservasiEmojiIndikatorRisiko:
    """Feature: sow-analysis-docx-conversion, Property 6: Preservasi Emoji Indikator Risiko

    For any konten Markdown yang mengandung emoji indikator risiko (🔴, 🟠, 🟡, 🟢),
    Dokumen_DOCX yang dihasilkan harus mengandung representasi yang setara untuk
    setiap emoji tersebut (baik emoji asli maupun teks berwarna pengganti).

    **Validates: Requirements 3.1**
    """

    # Mapping emoji → fallback text
    EMOJI_FALLBACK = {
        '🔴': '[CRITICAL]',
        '🟠': '[HIGH]',
        '🟡': '[MEDIUM]',
        '🟢': '[LOW]',
    }

    @settings(max_examples=100)
    @given(
        emojis=st.lists(
            st.sampled_from(['🔴', '🟠', '🟡', '🟢']),
            min_size=1,
            max_size=6,
        ),
        surrounding_text=st.text(
            alphabet=st.characters(
                whitelist_categories=("L", "N", "Zs"),
                blacklist_characters="\n\r\x00#*|`-[]",
            ),
            min_size=1,
            max_size=20,
        ),
    )
    def test_emoji_risiko_dipertahankan_di_docx(self, emojis, surrounding_text):
        """Generate random teks dengan emoji 🔴🟠🟡🟢, verifikasi representasi
        setara ada di DOCX (emoji asli atau fallback text).

        **Validates: Requirements 3.1**
        """
        clean_text = surrounding_text.strip()
        if not clean_text:
            clean_text = "Risk"

        # Build Markdown with heading and emoji content
        emoji_str = " ".join(emojis)
        md_content = f"# Report\n{clean_text} {emoji_str}"

        # Parse and build DOCX
        sections = parse_markdown(md_content)
        doc = build_docx(sections, title="Test Report")

        # Save to temp file and read back
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = tmp.name
            doc.save(tmp_path)

        try:
            read_doc = DocxDocument(tmp_path)

            # Collect all text from all paragraphs in the DOCX
            all_text = ""
            for para in read_doc.paragraphs:
                all_text += para.text + "\n"

            # For each emoji in input, verify either the emoji itself
            # or its fallback text is present in the DOCX
            for emoji in emojis:
                fallback = self.EMOJI_FALLBACK[emoji]
                assert emoji in all_text or fallback in all_text, (
                    f"Emoji '{emoji}' or fallback '{fallback}' not found in DOCX.\n"
                    f"DOCX text: {all_text!r}"
                )
        finally:
            os.unlink(tmp_path)


class TestPropertyPreservasiFormattingInlineDanList:
    """Feature: sow-analysis-docx-conversion, Property 7: Preservasi Formatting Inline dan List

    For any konten Markdown yang mengandung teks bold, italic, bullet list, atau
    checklist, Dokumen_DOCX yang dihasilkan harus mempertahankan formatting
    bold/italic pada teks yang sesuai, mengonversi bullet list menjadi list Word,
    dan mengonversi checklist menjadi representasi visual (☐/☑).

    **Validates: Requirements 2.3, 3.2, 3.4**
    """

    # Strategy: safe alphanumeric text without special Markdown characters
    _safe_text = st.text(
        alphabet=st.characters(
            whitelist_categories=("L", "N"),
            min_codepoint=65,
            max_codepoint=122,
        ),
        min_size=1,
        max_size=20,
    )

    @settings(max_examples=100)
    @given(
        bold_texts=st.lists(_safe_text, min_size=0, max_size=3),
        italic_texts=st.lists(_safe_text, min_size=0, max_size=3),
        bullet_texts=st.lists(_safe_text, min_size=0, max_size=3),
        checklist_items=st.lists(
            st.tuples(st.booleans(), _safe_text),
            min_size=0,
            max_size=3,
        ),
    )
    def test_formatting_inline_dan_list_dipertahankan(
        self, bold_texts, italic_texts, bullet_texts, checklist_items
    ):
        """Generate random bold/italic/bullet/checklist, verifikasi formatting di DOCX.

        **Validates: Requirements 2.3, 3.2, 3.4**
        """
        # Skip if nothing to test
        if not bold_texts and not italic_texts and not bullet_texts and not checklist_items:
            return

        # Build Markdown content
        md_lines = ["# TestDoc"]

        for txt in bold_texts:
            md_lines.append(f"This has **{txt}** word")

        for txt in italic_texts:
            md_lines.append(f"This has *{txt}* word")

        for txt in bullet_texts:
            md_lines.append(f"- {txt}")

        for checked, txt in checklist_items:
            mark = "x" if checked else " "
            md_lines.append(f"- [{mark}] {txt}")

        md_content = "\n".join(md_lines)

        # Parse and build DOCX
        sections = parse_markdown(md_content)
        doc = build_docx(sections, title="Test Formatting")

        # Save to temp file and read back
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = tmp.name
            doc.save(tmp_path)

        try:
            read_doc = DocxDocument(tmp_path)

            # Collect all paragraph text and run info
            all_text = ""
            bold_runs_found = []
            italic_runs_found = []
            for para in read_doc.paragraphs:
                all_text += para.text + "\n"
                for run in para.runs:
                    if run.bold:
                        bold_runs_found.append(run.text)
                    if run.italic:
                        italic_runs_found.append(run.text)

            # Verify bold text has bold formatting in DOCX runs
            for txt in bold_texts:
                assert any(txt in r for r in bold_runs_found), (
                    f"Bold text '{txt}' not found in bold runs.\n"
                    f"Bold runs: {bold_runs_found}\n"
                    f"All text: {all_text!r}"
                )

            # Verify italic text has italic formatting in DOCX runs
            for txt in italic_texts:
                assert any(txt in r for r in italic_runs_found), (
                    f"Italic text '{txt}' not found in italic runs.\n"
                    f"Italic runs: {italic_runs_found}\n"
                    f"All text: {all_text!r}"
                )

            # Verify bullet items appear in the DOCX
            for txt in bullet_texts:
                assert txt in all_text, (
                    f"Bullet item '{txt}' not found in DOCX text.\n"
                    f"All text: {all_text!r}"
                )

            # Verify checklist items have ☐ or ☑ prefix in the DOCX
            for checked, txt in checklist_items:
                expected_prefix = "☑" if checked else "☐"
                assert txt in all_text, (
                    f"Checklist item '{txt}' not found in DOCX text.\n"
                    f"All text: {all_text!r}"
                )
                assert expected_prefix in all_text, (
                    f"Checklist prefix '{expected_prefix}' not found in DOCX text.\n"
                    f"All text: {all_text!r}"
                )
        finally:
            os.unlink(tmp_path)


class TestPropertyFormattingCodeBlock:
    """Feature: sow-analysis-docx-conversion, Property 8: Formatting Code Block

    For any konten Markdown yang mengandung code block (dibatasi oleh triple backtick),
    Dokumen_DOCX yang dihasilkan harus mengandung paragraf dengan font monospace
    yang memuat konten code block tersebut.

    **Validates: Requirements 3.3**
    """

    @settings(max_examples=100)
    @given(
        code_content=st.text(
            alphabet=st.characters(
                whitelist_categories=("L", "N", "Zs"),
                blacklist_characters="\x00`",
            ),
            min_size=1,
            max_size=80,
        ),
    )
    def test_code_block_menggunakan_font_monospace(self, code_content):
        """Generate random code block content, verifikasi font monospace di DOCX.

        **Validates: Requirements 3.3**
        """
        clean_content = code_content.strip()
        if not clean_content:
            clean_content = "code"

        # Build Markdown with a heading and a code block
        md_content = f"# TestDoc\n```\n{clean_content}\n```"

        # Parse and build DOCX
        sections = parse_markdown(md_content)
        doc = build_docx(sections, title="Test Code Block")

        # Save to temp file and read back
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = tmp.name
            doc.save(tmp_path)

        try:
            read_doc = DocxDocument(tmp_path)

            # Find paragraphs containing the code block content,
            # excluding heading paragraphs to avoid false matches
            code_paragraphs = [
                para for para in read_doc.paragraphs
                if clean_content in para.text
                and not para.style.name.startswith("Heading")
            ]

            assert len(code_paragraphs) > 0, (
                f"Code block content '{clean_content}' not found in any non-heading DOCX paragraph."
            )

            # Verify font name is 'Courier New' for runs containing the code content
            for para in code_paragraphs:
                for run in para.runs:
                    if clean_content in run.text:
                        assert run.font.name == "Courier New", (
                            f"Expected font 'Courier New' for code block run, "
                            f"got '{run.font.name}'. Run text: '{run.text}'"
                        )
        finally:
            os.unlink(tmp_path)


# ---------------------------------------------------------------------------
# DOCX Builder Unit Tests
# ---------------------------------------------------------------------------


class TestBuildDocx:
    """Unit tests untuk build_docx() — verifikasi formatting DOCX output.

    Requirements: 2.1, 2.2, 2.3, 2.4, 2.5, 2.6, 3.1, 3.2, 3.3
    """

    def _build_and_read(self, sections, title="Test Document"):
        """Helper: build DOCX dari sections, simpan ke temp file, baca kembali."""
        doc = build_docx(sections, title=title)
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = tmp.name
            doc.save(tmp_path)
        return tmp_path

    # --- 1. Heading styles ---

    def test_heading_styles(self):
        """Heading H1, H2, H3 harus memiliki style Heading 1/2/3 yang benar."""
        sections = [
            Section(level=1, text="Heading One"),
            Section(level=2, text="Heading Two"),
            Section(level=3, text="Heading Three"),
        ]
        tmp_path = self._build_and_read(sections)
        try:
            read_doc = DocxDocument(tmp_path)
            heading_paras = [
                p for p in read_doc.paragraphs
                if p.style.name.startswith("Heading")
            ]
            assert len(heading_paras) == 3, (
                f"Expected 3 heading paragraphs, got {len(heading_paras)}"
            )
            assert heading_paras[0].style.name == "Heading 1"
            assert heading_paras[0].text == "Heading One"
            assert heading_paras[1].style.name == "Heading 2"
            assert heading_paras[1].text == "Heading Two"
            assert heading_paras[2].style.name == "Heading 3"
            assert heading_paras[2].text == "Heading Three"
        finally:
            os.unlink(tmp_path)

    # --- 2. Table header colored ---

    def test_table_header_colored(self):
        """Header row tabel harus memiliki background biru #4472C4 dan teks putih."""
        from docx.oxml.ns import qn

        table_obj = Table(headers=["Col A", "Col B"], rows=[["1", "2"]])
        sections = [Section(level=1, text="T", children=[table_obj])]
        tmp_path = self._build_and_read(sections)
        try:
            read_doc = DocxDocument(tmp_path)
            assert len(read_doc.tables) >= 1
            docx_table = read_doc.tables[0]
            header_row = docx_table.rows[0]
            for cell in header_row.cells:
                # Check shading fill color
                tc_pr = cell._tc.tcPr
                assert tc_pr is not None
                shading = tc_pr.find(qn('w:shd'))
                assert shading is not None
                fill = shading.get(qn('w:fill'))
                assert fill == '4472C4', f"Expected fill '4472C4', got '{fill}'"
                # Check text is white and bold (skip empty runs)
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text.strip():
                            assert run.bold is True
                            assert run.font.color.rgb is not None
                            assert str(run.font.color.rgb) == "FFFFFF"
        finally:
            os.unlink(tmp_path)

    # --- 3. Table alternating rows ---

    def test_table_alternating_rows(self):
        """Baris data tabel harus memiliki alternating row colors."""
        from docx.oxml.ns import qn

        table_obj = Table(
            headers=["H1", "H2"],
            rows=[["a", "b"], ["c", "d"], ["e", "f"], ["g", "h"]],
        )
        sections = [Section(level=1, text="T", children=[table_obj])]
        tmp_path = self._build_and_read(sections)
        try:
            read_doc = DocxDocument(tmp_path)
            docx_table = read_doc.tables[0]
            # Data rows start at index 1 (index 0 is header)
            for row_idx in range(1, len(docx_table.rows)):
                cell = docx_table.rows[row_idx].cells[0]
                tc_pr = cell._tc.tcPr
                shading = tc_pr.find(qn('w:shd')) if tc_pr is not None else None
                data_row_idx = row_idx - 1  # 0-based data row index
                if data_row_idx % 2 == 1:
                    # Odd data rows (1, 3, ...) should have gray background
                    assert shading is not None, f"Data row {data_row_idx} should have shading"
                    fill = shading.get(qn('w:fill'))
                    assert fill == 'F2F2F2', f"Expected fill 'F2F2F2' for row {data_row_idx}, got '{fill}'"
        finally:
            os.unlink(tmp_path)

    # --- 4. Checklist prefix ---

    def test_checklist_prefix(self):
        """ChecklistItem harus menggunakan prefix ☐ (unchecked) dan ☑ (checked)."""
        children = [
            ChecklistItem(runs=[TextRun(text="todo item")], checked=False),
            ChecklistItem(runs=[TextRun(text="done item")], checked=True),
        ]
        sections = [Section(level=1, text="CL", children=children)]
        tmp_path = self._build_and_read(sections)
        try:
            read_doc = DocxDocument(tmp_path)
            all_text = "\n".join(p.text for p in read_doc.paragraphs)
            assert "☐" in all_text, "Unchecked prefix ☐ not found"
            assert "todo item" in all_text
            assert "☑" in all_text, "Checked prefix ☑ not found"
            assert "done item" in all_text
        finally:
            os.unlink(tmp_path)

    # --- 5. Code block font ---

    def test_code_block_font(self):
        """CodeBlock harus menggunakan font Courier New."""
        children = [CodeBlock(content="print('hello')")]
        sections = [Section(level=1, text="CB", children=children)]
        tmp_path = self._build_and_read(sections)
        try:
            read_doc = DocxDocument(tmp_path)
            code_paras = [
                p for p in read_doc.paragraphs
                if "print('hello')" in p.text
            ]
            assert len(code_paras) >= 1, "Code block paragraph not found"
            for para in code_paras:
                for run in para.runs:
                    if "print('hello')" in run.text:
                        assert run.font.name == "Courier New", (
                            f"Expected 'Courier New', got '{run.font.name}'"
                        )
        finally:
            os.unlink(tmp_path)

    # --- 6. Header contains title ---

    def test_header_contains_title(self):
        """Header halaman harus berisi judul dokumen (rata kanan)."""
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        sections = [Section(level=1, text="Intro")]
        tmp_path = self._build_and_read(sections, title="My Report Title")
        try:
            read_doc = DocxDocument(tmp_path)
            header = read_doc.sections[0].header
            header_text = "".join(p.text for p in header.paragraphs)
            assert "My Report Title" in header_text, (
                f"Title not found in header. Header text: '{header_text}'"
            )
            # Verify right alignment
            for para in header.paragraphs:
                if "My Report Title" in para.text:
                    assert para.alignment == WD_ALIGN_PARAGRAPH.RIGHT
        finally:
            os.unlink(tmp_path)

    # --- 7. Footer has page number ---

    def test_footer_has_page_number(self):
        """Footer halaman harus memiliki field PAGE untuk nomor halaman."""
        from docx.oxml.ns import qn

        sections = [Section(level=1, text="Intro")]
        tmp_path = self._build_and_read(sections)
        try:
            read_doc = DocxDocument(tmp_path)
            footer = read_doc.sections[0].footer
            # Look for fldChar with PAGE instrText in footer XML
            footer_xml = footer._element.xml
            assert "PAGE" in footer_xml, (
                "PAGE field not found in footer XML"
            )
        finally:
            os.unlink(tmp_path)

    # --- 8. Page margins ---

    def test_page_margins(self):
        """Margin halaman harus 1 inci di semua sisi."""
        from docx.shared import Inches

        sections = [Section(level=1, text="Intro")]
        tmp_path = self._build_and_read(sections)
        try:
            read_doc = DocxDocument(tmp_path)
            section = read_doc.sections[0]
            one_inch = Inches(1)
            assert section.top_margin == one_inch, (
                f"Top margin: expected {one_inch}, got {section.top_margin}"
            )
            assert section.bottom_margin == one_inch, (
                f"Bottom margin: expected {one_inch}, got {section.bottom_margin}"
            )
            assert section.left_margin == one_inch, (
                f"Left margin: expected {one_inch}, got {section.left_margin}"
            )
            assert section.right_margin == one_inch, (
                f"Right margin: expected {one_inch}, got {section.right_margin}"
            )
        finally:
            os.unlink(tmp_path)

    # --- 9. Page size US Letter ---

    def test_page_size_us_letter(self):
        """Ukuran halaman harus US Letter (8.5" x 11")."""
        from docx.shared import Inches

        sections = [Section(level=1, text="Intro")]
        tmp_path = self._build_and_read(sections)
        try:
            read_doc = DocxDocument(tmp_path)
            section = read_doc.sections[0]
            assert section.page_width == Inches(8.5), (
                f"Page width: expected {Inches(8.5)}, got {section.page_width}"
            )
            assert section.page_height == Inches(11), (
                f"Page height: expected {Inches(11)}, got {section.page_height}"
            )
        finally:
            os.unlink(tmp_path)


# ---------------------------------------------------------------------------
# Property Test: Derivasi Path Output
# ---------------------------------------------------------------------------

from convert_md_to_docx import convert_md_to_docx


class TestPropertyDerivasiPathOutput:
    """Feature: sow-analysis-docx-conversion, Property 1: Derivasi Path Output

    For any file Markdown dengan path valid, path output DOCX yang dihasilkan
    harus berada di direktori yang sama dengan file input dan memiliki nama file
    yang sama dengan ekstensi `.docx` menggantikan `.md`.

    **Validates: Requirements 1.2, 1.3**
    """

    @settings(max_examples=100)
    @given(
        filename=st.text(
            alphabet=st.characters(
                whitelist_categories=("L", "N"),
                min_codepoint=65,
                max_codepoint=122,
            ),
            min_size=1,
            max_size=20,
        ),
    )
    def test_output_path_sama_dir_dan_nama_dengan_ekstensi_docx(self, filename):
        """Generate random filenames .md di random directories, verifikasi
        output path = same dir + same name + .docx.

        **Validates: Requirements 1.2, 1.3**
        """
        # Create a temp directory to simulate a random directory
        tmp_dir = tempfile.mkdtemp()
        try:
            md_filename = f"{filename}.md"
            md_path = os.path.join(tmp_dir, md_filename)

            # Create actual temp MD file with simple content
            with open(md_path, "w", encoding="utf-8") as f:
                f.write("# Test\nHello")

            # Call convert_md_to_docx without specifying output path
            success, message = convert_md_to_docx(md_path)

            # Derive expected output path
            expected_docx_path = os.path.join(tmp_dir, f"{filename}.docx")

            # Verify the output DOCX file was created in the same directory
            assert os.path.exists(expected_docx_path), (
                f"Expected DOCX at '{expected_docx_path}' but file does not exist.\n"
                f"Success: {success}, Message: {message}"
            )

            # Verify the output filename is the same as input but with .docx extension
            actual_docx_name = os.path.basename(expected_docx_path)
            expected_docx_name = f"{filename}.docx"
            assert actual_docx_name == expected_docx_name, (
                f"Expected filename '{expected_docx_name}', got '{actual_docx_name}'"
            )

            # Verify conversion was successful
            assert success is True, (
                f"Conversion failed: {message}"
            )
        finally:
            # Clean up temp files
            import shutil
            shutil.rmtree(tmp_dir, ignore_errors=True)


# ---------------------------------------------------------------------------
# Property Test: Kegagalan Non-Destruktif
# ---------------------------------------------------------------------------

import hashlib


class TestPropertyKegagalanNonDestruktif:
    """Feature: sow-analysis-docx-conversion, Property 9: Kegagalan Non-Destruktif

    For any skenario di mana konversi gagal (file corrupt, error parsing, dll),
    file Markdown sumber harus tetap tidak berubah dan fungsi konversi harus
    mengembalikan error code tanpa melempar exception yang tidak tertangani.

    **Validates: Requirements 4.3, 4.4**
    """

    @settings(max_examples=100)
    @given(
        md_content=st.text(
            alphabet=st.characters(
                whitelist_categories=("L", "N", "Zs"),
                blacklist_characters="\x00",
            ),
            min_size=1,
            max_size=200,
        ),
        heading=st.text(
            alphabet=st.characters(
                whitelist_categories=("L", "N"),
                min_codepoint=65,
                max_codepoint=122,
            ),
            min_size=1,
            max_size=30,
        ),
    )
    def test_kegagalan_tidak_mengubah_file_md_sumber(self, md_content, heading):
        """Generate random MD, simulasi error (invalid output path),
        verifikasi MD tidak berubah dan tidak ada unhandled exception.

        **Validates: Requirements 4.3, 4.4**
        """
        # Build a valid MD content with heading and body
        full_md = f"# {heading.strip() or 'Heading'}\n{md_content}"

        # Create a temp MD file with the generated content
        tmp_dir = tempfile.mkdtemp()
        try:
            md_path = os.path.join(tmp_dir, "test_input.md")
            with open(md_path, "w", encoding="utf-8") as f:
                f.write(full_md)

            # Record the MD file content hash before conversion
            with open(md_path, "rb") as f:
                hash_before = hashlib.sha256(f.read()).hexdigest()

            # Attempt conversion with an invalid/non-writable output path
            # Use a path to a non-existent directory so the save will fail
            invalid_output_path = os.path.join(
                tmp_dir, "nonexistent", "subdir", "output.docx"
            )

            # This should NOT raise an unhandled exception
            success, message = convert_md_to_docx(md_path, invalid_output_path)

            # Verify: function returns (False, error_message)
            assert success is False, (
                f"Expected conversion to fail with invalid output path, "
                f"but got success=True. Message: {message}"
            )
            assert isinstance(message, str) and len(message) > 0, (
                "Expected a non-empty error message on failure"
            )

            # Verify: the original MD file content is unchanged (hash matches)
            with open(md_path, "rb") as f:
                hash_after = hashlib.sha256(f.read()).hexdigest()

            assert hash_before == hash_after, (
                f"MD file was modified during failed conversion!\n"
                f"Hash before: {hash_before}\n"
                f"Hash after:  {hash_after}"
            )
        finally:
            import shutil
            shutil.rmtree(tmp_dir, ignore_errors=True)


# ---------------------------------------------------------------------------
# Error Handling Unit Tests
# ---------------------------------------------------------------------------

from unittest.mock import patch


class TestErrorHandling:
    """Unit tests untuk error handling convert_md_to_docx().

    Requirements: 4.1, 4.2, 4.3, 4.4
    """

    def test_file_not_found(self):
        """File tidak ditemukan menampilkan pesan error dengan path.

        Validates: Requirements 4.1
        """
        success, message = convert_md_to_docx("nonexistent_file.md")
        assert success is False
        assert "nonexistent_file.md" in message

    def test_dependency_missing(self):
        """Dependency missing menampilkan instruksi instalasi.

        Validates: Requirements 4.2
        """
        # Create a valid temp MD file
        tmp_dir = tempfile.mkdtemp()
        try:
            md_path = os.path.join(tmp_dir, "test.md")
            with open(md_path, "w", encoding="utf-8") as f:
                f.write("# Test\nHello world")

            # Mock the import of docx to raise ImportError
            import builtins
            original_import = builtins.__import__

            def mock_import(name, *args, **kwargs):
                if name == "docx":
                    raise ImportError("No module named 'docx'")
                return original_import(name, *args, **kwargs)

            with patch("builtins.__import__", side_effect=mock_import):
                success, message = convert_md_to_docx(md_path)

            assert success is False
            assert "pip install python-docx" in message
        finally:
            import shutil
            shutil.rmtree(tmp_dir, ignore_errors=True)

    def test_failed_conversion_preserves_md(self):
        """Konversi gagal tidak mengubah file MD sumber.

        Validates: Requirements 4.3
        """
        tmp_dir = tempfile.mkdtemp()
        try:
            md_path = os.path.join(tmp_dir, "source.md")
            original_content = "# Report\n\nThis is the original content.\n\n- Item 1\n- Item 2"
            with open(md_path, "w", encoding="utf-8") as f:
                f.write(original_content)

            # Use an invalid output path (non-existent nested directory)
            invalid_output = os.path.join(tmp_dir, "no", "such", "dir", "out.docx")
            success, message = convert_md_to_docx(md_path, invalid_output)

            assert success is False

            # Verify the MD file content is unchanged
            with open(md_path, "r", encoding="utf-8") as f:
                content_after = f.read()
            assert content_after == original_content
        finally:
            import shutil
            shutil.rmtree(tmp_dir, ignore_errors=True)

    def test_custom_output_path(self):
        """Output path custom dengan --output flag.

        Validates: Requirements 4.4
        """
        tmp_dir = tempfile.mkdtemp()
        try:
            md_path = os.path.join(tmp_dir, "input.md")
            with open(md_path, "w", encoding="utf-8") as f:
                f.write("# Custom Output Test\nSome content here.")

            custom_output = os.path.join(tmp_dir, "custom_name.docx")
            success, message = convert_md_to_docx(md_path, custom_output)

            assert success is True
            assert os.path.exists(custom_output)
        finally:
            import shutil
            shutil.rmtree(tmp_dir, ignore_errors=True)


# ---------------------------------------------------------------------------
# Property Test: Round-Trip Fidelitas Konten
# ---------------------------------------------------------------------------


class TestPropertyRoundTripFidelitasKonten:
    """Feature: sow-analysis-docx-conversion, Property 2: Round-Trip Fidelitas Konten

    For any file Laporan_Analisis Markdown yang valid, semua teks, data tabel,
    dan item daftar yang ada dalam file MD harus dapat ditemukan dalam
    Dokumen_DOCX yang dihasilkan tanpa ada konten yang hilang.

    **Validates: Requirements 7.1**
    """

    # Strategy: safe alphanumeric text (A-Z, a-z, 0-9 only)
    _safe_text = st.text(
        alphabet=st.characters(
            whitelist_categories=("L", "N"),
            min_codepoint=48,
            max_codepoint=122,
            blacklist_characters="[\\]^_`",
        ),
        min_size=1,
        max_size=20,
    )

    @settings(max_examples=100)
    @given(
        heading_text=st.text(
            alphabet=st.characters(
                whitelist_categories=("L", "N"),
                min_codepoint=65,
                max_codepoint=122,
            ),
            min_size=1,
            max_size=20,
        ),
        paragraphs=st.lists(
            st.text(
                alphabet=st.characters(
                    whitelist_categories=("L", "N"),
                    min_codepoint=65,
                    max_codepoint=122,
                ),
                min_size=1,
                max_size=30,
            ),
            min_size=1,
            max_size=3,
        ),
        table_data=st.data(),
        bullet_items=st.lists(
            st.text(
                alphabet=st.characters(
                    whitelist_categories=("L", "N"),
                    min_codepoint=65,
                    max_codepoint=122,
                ),
                min_size=1,
                max_size=20,
            ),
            min_size=1,
            max_size=3,
        ),
    )
    def test_round_trip_semua_konten_ditemukan_di_docx(
        self, heading_text, paragraphs, table_data, bullet_items
    ):
        """Generate random MD dengan heading, teks, tabel, list; extract teks
        dari DOCX dan bandingkan dengan MD source.

        **Validates: Requirements 7.1**
        """
        # --- Generate table dimensions and data ---
        num_cols = table_data.draw(st.integers(min_value=2, max_value=3))
        num_rows = table_data.draw(st.integers(min_value=1, max_value=3))

        cell_strategy = st.text(
            alphabet=st.characters(
                whitelist_categories=("L", "N"),
                min_codepoint=65,
                max_codepoint=122,
            ),
            min_size=1,
            max_size=10,
        )

        table_headers = [table_data.draw(cell_strategy) for _ in range(num_cols)]
        table_rows = []
        for _ in range(num_rows):
            row = [table_data.draw(cell_strategy) for _ in range(num_cols)]
            table_rows.append(row)

        # --- Build Markdown string ---
        md_lines = []

        # H1 heading
        md_lines.append(f"# {heading_text}")
        md_lines.append("")

        # Plain text paragraphs
        for para_text in paragraphs:
            md_lines.append(para_text)
            md_lines.append("")

        # Table
        header_line = "| " + " | ".join(table_headers) + " |"
        separator_line = "| " + " | ".join(["---"] * num_cols) + " |"
        md_lines.append(header_line)
        md_lines.append(separator_line)
        for row in table_rows:
            md_lines.append("| " + " | ".join(row) + " |")
        md_lines.append("")

        # Bullet items
        for item_text in bullet_items:
            md_lines.append(f"- {item_text}")

        md_content = "\n".join(md_lines)

        # --- Parse, build DOCX, save to temp file ---
        sections = parse_markdown(md_content)
        doc = build_docx(sections, title="Round Trip Test")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = tmp.name
            doc.save(tmp_path)

        try:
            # --- Read DOCX back ---
            read_doc = DocxDocument(tmp_path)

            # Extract all text from paragraphs
            all_para_text = "\n".join(p.text for p in read_doc.paragraphs)

            # Extract all text from table cells
            all_table_text = ""
            for tbl in read_doc.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        all_table_text += cell.text + "\n"

            # Combined text from the entire DOCX
            all_docx_text = all_para_text + "\n" + all_table_text

            # --- Verify heading text is present ---
            assert heading_text in all_docx_text, (
                f"Heading text '{heading_text}' not found in DOCX.\n"
                f"DOCX text: {all_docx_text!r}"
            )

            # --- Verify paragraph text is present ---
            for para_text in paragraphs:
                assert para_text in all_docx_text, (
                    f"Paragraph text '{para_text}' not found in DOCX.\n"
                    f"DOCX text: {all_docx_text!r}"
                )

            # --- Verify table header data is present ---
            for header in table_headers:
                assert header in all_docx_text, (
                    f"Table header '{header}' not found in DOCX.\n"
                    f"DOCX text: {all_docx_text!r}"
                )

            # --- Verify table cell data is present ---
            for row_idx, row in enumerate(table_rows):
                for col_idx, cell_val in enumerate(row):
                    assert cell_val in all_docx_text, (
                        f"Table cell [{row_idx}][{col_idx}] = '{cell_val}' "
                        f"not found in DOCX.\n"
                        f"DOCX text: {all_docx_text!r}"
                    )

            # --- Verify bullet item text is present ---
            for item_text in bullet_items:
                assert item_text in all_docx_text, (
                    f"Bullet item '{item_text}' not found in DOCX.\n"
                    f"DOCX text: {all_docx_text!r}"
                )
        finally:
            os.unlink(tmp_path)


# ---------------------------------------------------------------------------
# Integration Test End-to-End
# ---------------------------------------------------------------------------


class TestIntegrationEndToEnd:
    """Integration tests end-to-end untuk convert_md_to_docx.

    Requirements: 7.1, 7.2, 7.3
    """

    def test_konversi_sample_analysis_output(self):
        """Konversi sample-analysis-output.md menghasilkan DOCX valid.

        Validates: Requirements 7.1, 7.2, 7.3
        """
        # Construct path relative to this test file
        test_dir = os.path.dirname(__file__)
        sample_md = os.path.join(test_dir, "..", "examples", "sample-analysis-output.md")
        sample_md = os.path.normpath(sample_md)

        assert os.path.exists(sample_md), f"Sample file not found: {sample_md}"

        # Use a temp output path to avoid polluting the examples directory
        tmp_dir = tempfile.mkdtemp()
        try:
            output_docx = os.path.join(tmp_dir, "sample-analysis-output.docx")
            success, message = convert_md_to_docx(sample_md, output_docx)

            assert success is True, f"Conversion failed: {message}"
            assert os.path.exists(output_docx), "DOCX file was not created"

            # Verify DOCX is valid by opening with python-docx
            doc = DocxDocument(output_docx)
            assert len(doc.paragraphs) > 0, "DOCX has no paragraphs"
        finally:
            import shutil
            shutil.rmtree(tmp_dir, ignore_errors=True)

    def test_konversi_file_md_kosong(self):
        """Konversi file MD kosong tidak crash, returns tuple.

        Validates: Requirements 7.1
        """
        tmp_dir = tempfile.mkdtemp()
        try:
            md_path = os.path.join(tmp_dir, "empty.md")
            with open(md_path, "w", encoding="utf-8") as f:
                f.write("")

            result = convert_md_to_docx(md_path)

            # Should return a tuple (bool, str) without raising an exception
            assert isinstance(result, tuple), f"Expected tuple, got {type(result)}"
            assert len(result) == 2, f"Expected 2-element tuple, got {len(result)}"
            assert isinstance(result[0], bool)
            assert isinstance(result[1], str)
        finally:
            import shutil
            shutil.rmtree(tmp_dir, ignore_errors=True)

    def test_konversi_file_tanpa_heading(self):
        """Konversi file MD tanpa heading tetap menghasilkan DOCX valid.

        Validates: Requirements 7.1, 7.2
        """
        tmp_dir = tempfile.mkdtemp()
        try:
            md_path = os.path.join(tmp_dir, "no_heading.md")
            with open(md_path, "w", encoding="utf-8") as f:
                f.write("Just some plain text without any heading.\nAnother line of text.")

            output_docx = os.path.join(tmp_dir, "no_heading.docx")
            success, message = convert_md_to_docx(md_path, output_docx)

            assert success is True, f"Conversion failed: {message}"
            assert os.path.exists(output_docx), "DOCX file was not created"

            # Verify DOCX is valid
            doc = DocxDocument(output_docx)
            all_text = "\n".join(p.text for p in doc.paragraphs)
            assert "plain text" in all_text, "Content not found in DOCX"
        finally:
            import shutil
            shutil.rmtree(tmp_dir, ignore_errors=True)

    def test_tabel_dengan_kolom_tidak_rata(self):
        """Tabel dengan kolom tidak rata (rows have different column counts) tidak crash.

        Validates: Requirements 7.3
        """
        tmp_dir = tempfile.mkdtemp()
        try:
            md_content = (
                "# Test Table\n\n"
                "| A | B | C |\n"
                "|---|---|---|\n"
                "| 1 | 2 | 3 |\n"
                "| 4 | 5 |\n"
                "| 6 | 7 | 8 | 9 |\n"
            )
            md_path = os.path.join(tmp_dir, "uneven_table.md")
            with open(md_path, "w", encoding="utf-8") as f:
                f.write(md_content)

            output_docx = os.path.join(tmp_dir, "uneven_table.docx")
            success, message = convert_md_to_docx(md_path, output_docx)

            assert success is True, f"Conversion failed: {message}"
            assert os.path.exists(output_docx), "DOCX file was not created"

            # Verify DOCX is valid
            doc = DocxDocument(output_docx)
            assert len(doc.tables) >= 1, "No tables found in DOCX"
        finally:
            import shutil
            shutil.rmtree(tmp_dir, ignore_errors=True)
