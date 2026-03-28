#!/usr/bin/env python3
"""
convert_md_to_docx.py - Konversi laporan analisis Markdown menjadi dokumen Word (.docx).

Skrip ini mengonversi file laporan analisis SOW dalam format Markdown
menjadi dokumen Word (.docx) dengan formatting profesional.

Usage:
    python convert_md_to_docx.py <input_md_file> [--output <output_docx_file>]

Examples:
    python convert_md_to_docx.py analysis-report.md
    python convert_md_to_docx.py analysis-report.md --output report.docx
"""

import argparse
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Union, Optional


# ---------------------------------------------------------------------------
# Model Data Intermediate
# ---------------------------------------------------------------------------

@dataclass
class TextRun:
    """Representasi satu run teks dengan formatting."""
    text: str
    bold: bool = False
    italic: bool = False
    is_risk_indicator: bool = False
    font_name: Optional[str] = None
    font_size: Optional[int] = None
    color: Optional[str] = None


@dataclass
class Paragraph:
    """Representasi paragraf yang terdiri dari satu atau lebih TextRun."""
    runs: List[TextRun] = field(default_factory=list)


@dataclass
class BulletItem:
    """Representasi item bullet list dengan dukungan nested level."""
    runs: List[TextRun] = field(default_factory=list)
    level: int = 0


@dataclass
class ChecklistItem:
    """Representasi item checklist (- [ ] atau - [x])."""
    runs: List[TextRun] = field(default_factory=list)
    checked: bool = False


@dataclass
class CodeBlock:
    """Representasi blok kode (triple backtick)."""
    content: str = ""


@dataclass
class Table:
    """Representasi tabel Markdown dengan header dan baris data."""
    headers: List[str] = field(default_factory=list)
    rows: List[List[str]] = field(default_factory=list)


@dataclass
class Section:
    """Representasi section dokumen dengan heading dan children elements."""
    level: int = 1  # 1=H1, 2=H2, 3=H3
    text: str = ""
    children: List[Union['Paragraph', 'BulletItem', 'ChecklistItem', 'CodeBlock', 'Table', 'Section']] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Inline Text Parser
# ---------------------------------------------------------------------------

import re

RISK_EMOJIS = {'🔴', '🟠', '🟡', '🟢'}

# Pattern for bold (**...**), italic (*...*), and risk emoji
_INLINE_RE = re.compile(
    r'(\*\*(.+?)\*\*)'   # group 1,2: bold
    r'|(\*(.+?)\*)'      # group 3,4: italic
    r'|([🔴🟠🟡🟢])'     # group 5: risk emoji
)


def parse_inline(text: str) -> List[TextRun]:
    """Parse inline formatting dalam teks menjadi list TextRun.

    Menangani:
    - **bold** -> TextRun(bold=True)
    - *italic* -> TextRun(italic=True)
    - 🔴🟠🟡🟢 -> TextRun(is_risk_indicator=True)
    - Teks biasa -> TextRun()
    """
    if not text:
        return []

    runs: List[TextRun] = []
    last_end = 0

    for m in _INLINE_RE.finditer(text):
        # Add plain text before this match
        if m.start() > last_end:
            plain = text[last_end:m.start()]
            if plain:
                runs.append(TextRun(text=plain))

        if m.group(2) is not None:
            # Bold
            runs.append(TextRun(text=m.group(2), bold=True))
        elif m.group(4) is not None:
            # Italic
            runs.append(TextRun(text=m.group(4), italic=True))
        elif m.group(5) is not None:
            # Risk emoji
            runs.append(TextRun(text=m.group(5), is_risk_indicator=True))

        last_end = m.end()

    # Remaining plain text
    if last_end < len(text):
        remaining = text[last_end:]
        if remaining:
            runs.append(TextRun(text=remaining))

    # If no matches at all, return the whole text as a single run
    if not runs and text:
        runs.append(TextRun(text=text))

    return runs


# ---------------------------------------------------------------------------
# Markdown Parser
# ---------------------------------------------------------------------------

def _is_table_separator(line: str) -> bool:
    """Check if a line is a Markdown table separator (e.g. |---|---|)."""
    stripped = line.strip()
    if not stripped.startswith('|') or not stripped.endswith('|'):
        return False
    cells = [c.strip() for c in stripped.strip('|').split('|')]
    return all(re.match(r'^:?-{1,}:?$', c) for c in cells if c)


def _parse_table_row(line: str) -> List[str]:
    """Parse a single table row into a list of cell strings."""
    stripped = line.strip()
    if stripped.startswith('|'):
        stripped = stripped[1:]
    if stripped.endswith('|'):
        stripped = stripped[:-1]
    return [cell.strip() for cell in stripped.split('|')]


def parse_markdown(content: str) -> List[Section]:
    """Parse konten Markdown menjadi list of Section objects.

    Menangani:
    - Heading level 1-3 (#, ##, ###)
    - Tabel Markdown (header + separator + rows)
    - Bullet list (- item) dengan nested level
    - Checklist (- [ ], - [x])
    - Code block (triple backtick)
    - Inline formatting (**bold**, *italic*)
    - Emoji indikator risiko (🔴, 🟠, 🟡, 🟢)
    - Teks biasa sebagai Paragraph
    - Horizontal rule (---) sebagai separator

    Returns:
        Flat list of Section objects. Content before the first heading
        goes into a Section with level=0.
    """
    sections: List[Section] = []
    current_section: Optional[Section] = None
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_block_content: List[str] = []

    while i < len(lines):
        line = lines[i]

        # --- Code block handling ---
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                in_code_block = False
                block = CodeBlock(content='\n'.join(code_block_content))
                if current_section is None:
                    current_section = Section(level=0, text="")
                    sections.append(current_section)
                current_section.children.append(block)
                code_block_content = []
            else:
                # Start code block
                in_code_block = True
                code_block_content = []
            i += 1
            continue

        if in_code_block:
            code_block_content.append(line)
            i += 1
            continue

        # --- Heading ---
        heading_match = re.match(r'^(#{1,3})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            current_section = Section(level=level, text=text)
            sections.append(current_section)
            i += 1
            continue

        # --- Horizontal rule ---
        if re.match(r'^-{3,}\s*$', line.strip()) and not line.strip().startswith('- '):
            # Horizontal rule — add as a Paragraph with "---"
            if current_section is None:
                current_section = Section(level=0, text="")
                sections.append(current_section)
            current_section.children.append(Paragraph(runs=[TextRun(text="---")]))
            i += 1
            continue

        # --- Table detection ---
        # A table starts with a row containing |, followed by a separator row
        if '|' in line and (i + 1) < len(lines) and _is_table_separator(lines[i + 1]):
            headers = _parse_table_row(line)
            i += 2  # skip header and separator
            rows: List[List[str]] = []
            while i < len(lines) and '|' in lines[i] and lines[i].strip():
                row_cells = _parse_table_row(lines[i])
                rows.append(row_cells)
                i += 1
            table = Table(headers=headers, rows=rows)
            if current_section is None:
                current_section = Section(level=0, text="")
                sections.append(current_section)
            current_section.children.append(table)
            continue

        # --- Checklist ---
        checklist_match = re.match(r'^(\s*)- \[([ xX])\]\s*(.*)$', line)
        if checklist_match:
            checked = checklist_match.group(2).lower() == 'x'
            item_text = checklist_match.group(3)
            runs = parse_inline(item_text)
            item = ChecklistItem(runs=runs, checked=checked)
            if current_section is None:
                current_section = Section(level=0, text="")
                sections.append(current_section)
            current_section.children.append(item)
            i += 1
            continue

        # --- Bullet list ---
        bullet_match = re.match(r'^(\s*)- (.*)$', line)
        if bullet_match:
            indent = len(bullet_match.group(1))
            level = indent // 2  # 2 spaces per indent level
            item_text = bullet_match.group(2)
            runs = parse_inline(item_text)
            item = BulletItem(runs=runs, level=level)
            if current_section is None:
                current_section = Section(level=0, text="")
                sections.append(current_section)
            current_section.children.append(item)
            i += 1
            continue

        # --- Empty line (skip) ---
        if not line.strip():
            i += 1
            continue

        # --- Plain text / paragraph ---
        runs = parse_inline(line)
        if runs:
            para = Paragraph(runs=runs)
            if current_section is None:
                current_section = Section(level=0, text="")
                sections.append(current_section)
            current_section.children.append(para)

        i += 1

    # Handle unclosed code block
    if in_code_block and code_block_content:
        block = CodeBlock(content='\n'.join(code_block_content))
        if current_section is None:
            current_section = Section(level=0, text="")
            sections.append(current_section)
        current_section.children.append(block)

    return sections


# ---------------------------------------------------------------------------
# DOCX Builder
# ---------------------------------------------------------------------------

# Risk emoji → fallback mapping
_RISK_EMOJI_FALLBACK = {
    '🔴': ('[CRITICAL]', 'FF0000'),
    '🟠': ('[HIGH]', 'FF8C00'),
    '🟡': ('[MEDIUM]', '808000'),
    '🟢': ('[LOW]', '008000'),
}


def _apply_text_run(docx_run, text_run: TextRun) -> None:
    """Apply TextRun formatting to a python-docx Run object."""
    from docx.shared import Pt, RGBColor

    docx_run.text = text_run.text
    if text_run.bold:
        docx_run.bold = True
    if text_run.italic:
        docx_run.italic = True
    if text_run.font_name:
        docx_run.font.name = text_run.font_name
    if text_run.font_size:
        docx_run.font.size = Pt(text_run.font_size)
    if text_run.color:
        docx_run.font.color.rgb = RGBColor.from_string(text_run.color)


def _render_risk_emoji(paragraph, text_run: TextRun) -> None:
    """Render a risk emoji TextRun: try emoji first, add colored fallback."""
    from docx.shared import RGBColor

    emoji = text_run.text
    # First run: the emoji itself
    run = paragraph.add_run(emoji)
    # Second run: colored fallback text
    if emoji in _RISK_EMOJI_FALLBACK:
        fallback_text, color_hex = _RISK_EMOJI_FALLBACK[emoji]
        fb_run = paragraph.add_run(fallback_text)
        fb_run.font.color.rgb = RGBColor.from_string(color_hex)


def _render_runs(paragraph, runs: List[TextRun]) -> None:
    """Render a list of TextRun objects into a python-docx paragraph."""
    from docx.shared import Pt, RGBColor

    for tr in runs:
        if tr.is_risk_indicator:
            _render_risk_emoji(paragraph, tr)
        else:
            run = paragraph.add_run(tr.text)
            if tr.bold:
                run.bold = True
            if tr.italic:
                run.italic = True
            if tr.font_name:
                run.font.name = tr.font_name
            if tr.font_size:
                run.font.size = Pt(tr.font_size)
            if tr.color:
                run.font.color.rgb = RGBColor.from_string(tr.color)


def _set_cell_shading(cell, color_hex: str) -> None:
    """Set background shading on a table cell."""
    from docx.oxml.ns import qn
    from lxml import etree

    tc_pr = cell._tc.get_or_add_tcPr()
    shading = etree.SubElement(tc_pr, qn('w:shd'))
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), color_hex)


def _set_table_borders(table) -> None:
    """Set borders on all cells of a table."""
    from docx.oxml.ns import qn
    from lxml import etree

    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else etree.SubElement(tbl, qn('w:tblPr'))
    borders = etree.SubElement(tbl_pr, qn('w:tblBorders'))
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        element = etree.SubElement(borders, qn(f'w:{edge}'))
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), '000000')


def _add_page_number_field(paragraph) -> None:
    """Add a Word PAGE field code to a paragraph for page numbering."""
    from docx.oxml.ns import qn
    from lxml import etree

    run = paragraph.add_run()
    fld_char_begin = etree.SubElement(run._r, qn('w:fldChar'))
    fld_char_begin.set(qn('w:fldCharType'), 'begin')

    run2 = paragraph.add_run()
    instr_text = etree.SubElement(run2._r, qn('w:instrText'))
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = ' PAGE '

    run3 = paragraph.add_run()
    fld_char_end = etree.SubElement(run3._r, qn('w:fldChar'))
    fld_char_end.set(qn('w:fldCharType'), 'end')


def build_docx(sections: List[Section], title: str):
    """Membangun dokumen DOCX dari list Section.

    Args:
        sections: List of Section objects dari parse_markdown().
        title: Judul dokumen untuk header halaman.

    Returns:
        python-docx Document object.
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn

    doc = Document()

    # --- Page setup: US Letter, 1 inch margins ---
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # --- Default font: Arial 11pt ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # --- Configure heading styles ---
    heading_config = {
        1: ('Heading 1', 16),
        2: ('Heading 2', 14),
        3: ('Heading 3', 12),
    }
    for level, (style_name, size) in heading_config.items():
        h_style = doc.styles[style_name]
        h_style.font.name = 'Arial'
        h_style.font.size = Pt(size)
        h_style.font.bold = True

    # --- Header: document title (right-aligned) ---
    header = doc.sections[0].header
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header_para.add_run(title)
    header_run.font.name = 'Arial'
    header_run.font.size = Pt(9)

    # --- Footer: page number (center-aligned) ---
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_number_field(footer_para)

    # --- Render sections ---
    for sect in sections:
        _render_section(doc, sect)

    return doc


def _render_section(doc, section: Section) -> None:
    """Render a single Section and its children into the document."""
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    # Render heading (skip level 0 — preamble sections)
    if section.level >= 1:
        style_name = f'Heading {section.level}'
        heading_para = doc.add_heading(section.text, level=section.level)

    # Render children
    for child in section.children:
        if isinstance(child, Table):
            _render_table(doc, child)
        elif isinstance(child, BulletItem):
            _render_bullet(doc, child)
        elif isinstance(child, ChecklistItem):
            _render_checklist(doc, child)
        elif isinstance(child, CodeBlock):
            _render_code_block(doc, child)
        elif isinstance(child, Paragraph):
            _render_paragraph(doc, child)


def _render_table(doc, table: Table) -> None:
    """Render a Table as a Word table with styled header and alternating rows."""
    from docx.shared import Pt, RGBColor

    num_cols = len(table.headers)
    num_rows = len(table.rows) + 1  # +1 for header row
    docx_table = doc.add_table(rows=num_rows, cols=num_cols)
    docx_table.style = 'Table Grid'
    _set_table_borders(docx_table)

    # Header row: blue background (#4472C4), white bold text
    for col_idx, header_text in enumerate(table.headers):
        cell = docx_table.rows[0].cells[col_idx]
        _set_cell_shading(cell, '4472C4')
        cell.text = ''
        para = cell.paragraphs[0]
        run = para.add_run(header_text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Arial'
        run.font.size = Pt(11)

    # Data rows: alternating row colors
    for row_idx, row_data in enumerate(table.rows):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < num_cols:
                cell = docx_table.rows[row_idx + 1].cells[col_idx]
                # Alternating: odd rows (0-indexed) get gray background
                if row_idx % 2 == 1:
                    _set_cell_shading(cell, 'F2F2F2')
                cell.text = ''
                para = cell.paragraphs[0]
                # Parse inline formatting in cell text
                runs = parse_inline(cell_text)
                _render_runs(para, runs)


def _render_bullet(doc, bullet: BulletItem) -> None:
    """Render a BulletItem as a Word bullet list paragraph with level support."""
    from docx.shared import Pt

    para = doc.add_paragraph(style='List Bullet')
    # Set indentation level via paragraph format
    if bullet.level > 0:
        para.paragraph_format.left_indent = Pt(36 * bullet.level)
    _render_runs(para, bullet.runs)


def _render_checklist(doc, item: ChecklistItem) -> None:
    """Render a ChecklistItem with ☐ or ☑ prefix."""
    prefix = '☑ ' if item.checked else '☐ '
    para = doc.add_paragraph()
    run = para.add_run(prefix)
    run.font.name = 'Arial'
    _render_runs(para, item.runs)


def _render_code_block(doc, block: CodeBlock) -> None:
    """Render a CodeBlock with Courier New 9pt and background #F5F5F5."""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from lxml import etree

    para = doc.add_paragraph()
    # Set paragraph background shading
    p_pr = para._p.get_or_add_pPr()
    shading = etree.SubElement(p_pr, qn('w:shd'))
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'F5F5F5')

    run = para.add_run(block.content)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)


def _render_paragraph(doc, paragraph: Paragraph) -> None:
    """Render a Paragraph with its TextRuns."""
    para = doc.add_paragraph()
    _render_runs(para, paragraph.runs)


# ---------------------------------------------------------------------------
# Fungsi Konversi Utama
# ---------------------------------------------------------------------------

def convert_md_to_docx(md_path: str, docx_path: str = None) -> tuple[bool, str]:
    """Konversi file Markdown menjadi dokumen Word (.docx).

    Args:
        md_path: Path ke file Markdown sumber.
        docx_path: Path output file DOCX (opsional). Jika None, akan diderivasi
                   dari md_path dengan ekstensi .docx.

    Returns:
        Tuple (success, message):
        - (True, pesan_sukses) jika konversi berhasil
        - (False, pesan_error) jika konversi gagal
    """
    try:
        import docx  # noqa: F401
    except ImportError:
        return False, "python-docx belum terinstal. Jalankan: pip install python-docx"

    md_file = Path(md_path)
    if not md_file.exists():
        return False, f"File tidak ditemukan: {md_path}"

    if docx_path is None:
        docx_path = str(md_file.with_suffix('.docx'))

    try:
        content = md_file.read_text(encoding='utf-8')
        sections = parse_markdown(content)
        title = sections[0].text if sections else "SOW Analysis Report"
        doc = build_docx(sections, title)
        doc.save(docx_path)
        return True, f"✅ Dokumen DOCX berhasil dibuat: {docx_path}"
    except Exception as e:
        return False, f"Konversi DOCX gagal: {str(e)}"


# ---------------------------------------------------------------------------
# CLI Entry Point
# ---------------------------------------------------------------------------

def main() -> int:
    """Entry point CLI untuk konversi MD ke DOCX.

    Returns:
        Exit code: 0 jika sukses, 1 jika gagal.
    """
    parser = argparse.ArgumentParser(
        description="Konversi laporan analisis Markdown (.md) menjadi dokumen Word (.docx).",
        usage="%(prog)s <input_md_file> [--output <output_docx_file>]",
    )
    parser.add_argument(
        "input_md_file",
        help="Path ke file Markdown yang akan dikonversi",
    )
    parser.add_argument(
        "--output",
        metavar="<output_docx_file>",
        default=None,
        help="Path output file DOCX (default: nama file yang sama dengan ekstensi .docx)",
    )

    args = parser.parse_args()

    success, message = convert_md_to_docx(args.input_md_file, args.output)
    print(message)
    return 0 if success else 1


if __name__ == "__main__":
    sys.exit(main())
